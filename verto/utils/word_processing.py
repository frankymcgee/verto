import re
import frappe
from docx import Document

def load_document(path):
    return Document(path)

def save_document(doc, path):
    doc.save(path)

def replace_merge_fields(doc, merge_fields):
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        for key, value in merge_fields.items():
            pattern = r'«' + key + r'»'
            if re.search(pattern, original_text):
                new_text = re.sub(pattern, str(value), original_text)
                paragraph.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                for key, value in merge_fields.items():
                    pattern = r'«' + key + r'»'
                    if re.search(pattern, original_text):
                        new_text = re.sub(pattern, str(value), original_text)
                        cell.text = new_text

def replace_image(doc, replacement_id, new_image_path):
    def log_images_in_part(part):
        image_count = 1
        for rel in part.rels.values():
            if "image" in rel.target_ref:
                image_name = rel.target_ref
                frappe.log_error(f"Found image {image_count}: Relationship ID: {rel.rId}, Target Ref: {image_name}", "Image Logging")
                image_count += 1

    def replace_in_part(part):
        for rel in part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_name = rel.target_ref
                    frappe.log_error(f"Checking image: Relationship ID: {rel.rId}, Target Ref: {image_name}", "Image Replacement")
                    # If the target reference contains the replacement ID, replace the image
                    if replacement_id in image_name:
                        with open(new_image_path, 'rb') as new_image_file:
                            rel.target_part._blob = new_image_file.read()
                        frappe.log_error(f"Replaced image with target ref containing: {replacement_id}", "Image Replacement")
                        return True
                except AttributeError:
                    pass
        return False

    # Log images in main document
    log_images_in_part(doc.part)
    if replace_in_part(doc.part):
        return
    
    # Log images in headers and footers
    for section in doc.sections:
        log_images_in_part(section.header.part)
        if replace_in_part(section.header.part):
            return
        log_images_in_part(section.footer.part)
        if replace_in_part(section.footer.part):
            return
    
    frappe.log_error(f"No image with target ref containing: {replacement_id} found", "Image Replacement")

def update_table_of_contents(doc):
    # This function requires Word to update the TOC; Python-docx does not support TOC updates
    for para in doc.paragraphs:
        if 'TOC' in para.style.name:
            toc = para
            break
    else:
        toc = None

    if toc is not None:
        toc._element.clear_content()
        toc_text = OxmlElement("w:fldSimple")
        toc_text.set(qn("w:instr"), "TOC \\o 1-3 \\h \\z \\u")
        toc.append(toc_text)

def process_document(input_path, output_path, merge_fields, image_replacements=None):
    doc = load_document(input_path)
    replace_merge_fields(doc, merge_fields)
    
    if image_replacements:
        for replacement_id, new_image_path in image_replacements.items():
            replace_image(doc, replacement_id, new_image_path)
    
    update_table_of_contents(doc)  # Placeholder for updating TOC
    
    save_document(doc, output_path)
